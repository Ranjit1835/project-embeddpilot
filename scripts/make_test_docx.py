"""Spike aid: build a small DOCX register map to exercise the DOCX path."""

import docx

ROWS = [
    ("Register Name", "Address Offset", "Reset Value", "Access", "Description"),
    ("CTRL", "0x00", "0x00", "R/W", "Control register"),
    ("STATUS", "0x04", "0x01", "RO", "Status register"),
    ("DATA", "0x08", "0x00", "R/W", "Data register"),
    ("IRQ_EN", "0x0C", "0x00", "R/W", "Interrupt enable"),
]

d = docx.Document()
d.add_heading("FAKE1234 Register Map", level=1)
table = d.add_table(rows=len(ROWS), cols=len(ROWS[0]))
for i, row in enumerate(ROWS):
    for j, cell in enumerate(row):
        table.cell(i, j).text = cell
d.save("extraction/input/fake_regmap.docx")
print("wrote extraction/input/fake_regmap.docx")
