"""Stage 1: load a PDF or DOCX into per-page text + raw tables.

Provenance rule: every downstream artifact must be traceable to page numbers,
so this is the only module that touches the source file.
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field

MAX_FILE_BYTES = 50 * 1024 * 1024  # 50MB per spec
# A page with images but almost no text layer is treated as scanned.
SCANNED_PAGE_MIN_CHARS = 40


@dataclass
class Page:
    number: int                       # 1-based; for DOCX this is a block index
    text: str
    tables: list[list[list[str]]] = field(default_factory=list)  # tables -> rows -> cells
    is_scanned: bool = False


@dataclass
class Document:
    path: str
    kind: str                         # "pdf" | "docx"
    pages: list[Page] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def scanned_pages(self) -> list[int]:
        return [p.number for p in self.pages if p.is_scanned]


class IngestionError(Exception):
    pass


def load_document(
    path: str,
    max_pages: int | None = None,
    page_range: tuple[int, int] | None = None,
) -> Document:
    if not os.path.isfile(path):
        raise IngestionError(f"file not found: {path}")
    size = os.path.getsize(path)
    if size > MAX_FILE_BYTES:
        raise IngestionError(
            f"file is {size / 1024 / 1024:.1f}MB, over the 50MB limit"
        )
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _load_pdf(path, max_pages, page_range)
    if ext == ".docx":
        return _load_docx(path)
    raise IngestionError(f"unsupported file type: {ext} (accepted: .pdf, .docx)")


def _clean_cell(cell) -> str:
    if cell is None:
        return ""
    # pdfplumber keeps intra-cell line breaks; collapse to single spaces
    return " ".join(str(cell).split())


def _load_pdf(
    path: str, max_pages: int | None, page_range: tuple[int, int] | None
) -> Document:
    import pdfplumber

    doc = Document(path=path, kind="pdf")
    with pdfplumber.open(path) as pdf:
        pages = pdf.pages if max_pages is None else pdf.pages[:max_pages]
        if page_range:
            # text/table extraction dominates runtime on 1000-page manuals;
            # skip pages the caller has already scoped out
            pages = [p for p in pages if page_range[0] <= p.page_number <= page_range[1]]
        for p in pages:
            text = p.extract_text() or ""
            raw_tables = p.extract_tables() or []
            tables = [
                [[_clean_cell(c) for c in row] for row in t if row]
                for t in raw_tables
            ]
            is_scanned = len(text.strip()) < SCANNED_PAGE_MIN_CHARS and bool(p.images)
            doc.pages.append(
                Page(number=p.page_number, text=text, tables=tables, is_scanned=is_scanned)
            )
            # release per-page layout caches; matters on 1000+ page manuals
            p.flush_cache()
    if doc.scanned_pages:
        doc.warnings.append(
            f"pages {doc.scanned_pages} appear scanned (image-only); "
            "text/table extraction skipped there — OCR fallback available at lower confidence"
        )
    return doc


def _load_docx(path: str) -> Document:
    """DOCX has no fixed pages; each top-level block (paragraph run or table)
    becomes a pseudo-page so provenance stays meaningful."""
    import docx

    d = docx.Document(path)
    doc = Document(path=path, kind="docx")
    doc.warnings.append("DOCX source: 'pages' are sequential block indices, not print pages")

    block_no = 0
    text_buf: list[str] = []
    for para in d.paragraphs:
        if para.text.strip():
            text_buf.append(para.text)
    block_no += 1
    doc.pages.append(Page(number=block_no, text="\n".join(text_buf)))

    for table in d.tables:
        block_no += 1
        rows = [[_clean_cell(c.text) for c in row.cells] for row in table.rows]
        doc.pages.append(Page(number=block_no, text="", tables=[rows]))
    return doc
